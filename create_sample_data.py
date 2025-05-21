from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def create_sample_docx():
    # Create a new Document
    doc = Document()
    
    # Add title
    title = doc.add_heading('PCOS Patient Data', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Read the sample data
    with open('data/raw/sample_patient_data.docx', 'r', encoding='utf-8') as f:
        content = f.read()
    
    # Split content into sections
    sections = content.split('\n\n')
    
    # Add each section to the document
    for section in sections:
        if section.strip():
            # Add section heading
            if ':' in section.split('\n')[0]:
                heading = section.split('\n')[0]
                doc.add_heading(heading, level=1)
                
                # Add section content
                content = '\n'.join(section.split('\n')[1:])
                paragraph = doc.add_paragraph(content)
            else:
                paragraph = doc.add_paragraph(section)
            
            # Set font size
            for run in paragraph.runs:
                run.font.size = Pt(11)
    
    # Save the document
    os.makedirs('data/raw', exist_ok=True)
    doc.save('data/raw/sample_patient_data.docx')
    print("Sample DOCX file created successfully!")

if __name__ == "__main__":
    create_sample_docx() 